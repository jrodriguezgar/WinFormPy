"""
Document Backend - Base classes for document handling

This module provides abstract base classes and implementations for
loading and rendering different document types.

Note: Uses lazy imports for all optional dependencies (PIL, PyMuPDF, python-docx).
"""

from abc import ABC, abstractmethod
from typing import List, Optional, Tuple, Any
import io
import sys
import subprocess


# =============================================================
# Lazy Library Import Management (supports pip and uv)
# =============================================================

def _is_uv_managed_environment() -> bool:
    """Check if the current Python environment is managed by uv."""
    if os.environ.get('UV_PROJECT_ENVIRONMENT'):
        return True
    if 'uv' in sys.executable.lower():
        return True
    exe_dir = os.path.dirname(sys.executable)
    venv_dir = os.path.dirname(exe_dir)
    pyvenv_cfg = os.path.join(venv_dir, 'pyvenv.cfg')
    if os.path.exists(pyvenv_cfg):
        try:
            with open(pyvenv_cfg, 'r') as f:
                content = f.read()
                if 'uv =' in content or 'uv=' in content:
                    return True
        except:
            pass
    return False


def _find_pyproject_dir() -> str:
    """Find the directory containing pyproject.toml."""
    current = os.getcwd()
    while current != os.path.dirname(current):
        if os.path.exists(os.path.join(current, 'pyproject.toml')):
            return current
        current = os.path.dirname(current)
    return None


def install_library(library_name: str, import_name: str = None) -> bool:
    """
    Checks if a library is installed and, if not, attempts to install it.
    Uses 'uv add' if the environment is uv-managed, otherwise uses pip.
    """
    check_name = import_name if import_name else library_name
    try:
        __import__(check_name)
        return True
    except ImportError:
        print(f"Installing '{library_name}'...")
        try:
            if _is_uv_managed_environment():
                project_dir = _find_pyproject_dir()
                if project_dir:
                    subprocess.check_call(["uv", "add", library_name], cwd=project_dir)
                else:
                    subprocess.check_call(["uv", "pip", "install", "--system", library_name])
            else:
                subprocess.check_call([sys.executable, "-m", "pip", "install", library_name])
            print(f"✓ '{library_name}' installed")
            return True
        except subprocess.CalledProcessError:
            print(f"✗ Failed to install '{library_name}'")
            return False
        except FileNotFoundError:
            try:
                subprocess.check_call([sys.executable, "-m", "pip", "install", library_name])
                return True
            except subprocess.CalledProcessError:
                print(f"✗ Failed. For uv envs: uv add {library_name}")
                return False


class DocumentBackend(ABC):
    """
    Abstract base class for document backends.
    
    Each document type (PDF, Word, Excel, etc.) should implement
    this interface to provide document loading and rendering.
    
    Note: PIL is loaded lazily when needed via install_library.
    """
    
    _Image = None
    _ImageDraw = None
    _ImageFont = None
    
    @classmethod
    def _ensure_pil(cls):
        """Lazy import of PIL library with auto-install."""
        if cls._Image is None:
            if install_library("Pillow", "PIL"):
                try:
                    from PIL import Image, ImageDraw, ImageFont
                    cls._Image = Image
                    cls._ImageDraw = ImageDraw
                    cls._ImageFont = ImageFont
                except ImportError:
                    raise ImportError(
                        "Pillow (PIL) is required for document viewing.\n"
                        "Install it with: pip install Pillow"
                    )
            else:
                raise ImportError(
                    "Pillow (PIL) is required for document viewing.\n"
                    "Install it with: pip install Pillow"
                )
        return cls._Image, cls._ImageDraw, cls._ImageFont
    
    def __init__(self):
        self.page_count = 0
        self.current_page = 0
        self.file_path = None
        self.document_loaded = False
    
    @abstractmethod
    def load_document(self, file_path: str) -> bool:
        """
        Load a document from file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            True if loaded successfully, False otherwise
        """
        pass
    
    @abstractmethod
    def get_page_count(self) -> int:
        """Get total number of pages in document."""
        pass
    
    @abstractmethod
    def render_page(self, page_number: int, width: int = None, height: int = None):
        """
        Render a specific page as PIL Image.
        
        Args:
            page_number: Page number (0-based)
            width: Target width (None for original)
            height: Target height (None for original)
            
        Returns:
            PIL Image of the rendered page
        """
        pass
    
    @abstractmethod
    def get_page_text(self, page_number: int) -> str:
        """
        Extract text from a specific page.
        
        Args:
            page_number: Page number (0-based)
            
        Returns:
            Text content of the page
        """
        pass
    
    def close_document(self):
        """Close the document and free resources."""
        self.document_loaded = False
        self.page_count = 0
        self.current_page = 0
    
    def get_document_info(self) -> dict:
        """
        Get document metadata.
        
        Returns:
            Dictionary with document information
        """
        return {
            'file_path': self.file_path,
            'page_count': self.page_count,
            'format': self.__class__.__name__.replace('Backend', '')
        }


class PDFBackend(DocumentBackend):
    """
    Backend for PDF document viewing.
    Requires: PyMuPDF (fitz) library
    
    Note: Uses lazy import - fitz is only imported when load_document is called.
    """
    
    def __init__(self):
        super().__init__()
        self.pdf_document = None
        self._fitz = None
    
    def _ensure_fitz(self):
        """Lazy import of PyMuPDF library with auto-install."""
        if self._fitz is None:
            if install_library("PyMuPDF", "fitz"):
                try:
                    import fitz  # PyMuPDF
                    self._fitz = fitz
                except ImportError:
                    raise ImportError(
                        "PyMuPDF (fitz) is required for PDF support.\n"
                        "Install it with: pip install PyMuPDF"
                    )
            else:
                raise ImportError(
                    "PyMuPDF (fitz) is required for PDF support.\n"
                    "Install it with: pip install PyMuPDF"
                )
        return self._fitz
    
    def load_document(self, file_path: str) -> bool:
        """Load PDF document."""
        try:
            fitz = self._ensure_fitz()
            self.pdf_document = fitz.open(file_path)
            self.file_path = file_path
            self.page_count = len(self.pdf_document)
            self.current_page = 0
            self.document_loaded = True
            return True
        except ImportError as e:
            raise e
        except Exception as e:
            print(f"Error loading PDF: {e}")
            return False
    
    def get_page_count(self) -> int:
        """Get total pages in PDF."""
        return self.page_count
    
    def render_page(self, page_number: int, width: int = None, height: int = None):
        """Render PDF page as image."""
        if not self.document_loaded or page_number >= self.page_count:
            return self._create_error_page("Invalid page number")
        
        try:
            Image, ImageDraw, ImageFont = self._ensure_pil()
            fitz = self._ensure_fitz()
            page = self.pdf_document[page_number]
            
            # Calculate zoom factor if width/height specified
            zoom = 1.0
            if width:
                page_width = page.rect.width
                zoom = width / page_width
            
            # Render page to pixmap
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat)
            
            # Convert to PIL Image
            img_data = pix.tobytes("ppm")
            img = Image.open(io.BytesIO(img_data))
            
            return img
        except Exception as e:
            return self._create_error_page(f"Error rendering page: {e}")
    
    def get_page_text(self, page_number: int) -> str:
        """Extract text from PDF page."""
        if not self.document_loaded or page_number >= self.page_count:
            return ""
        
        try:
            page = self.pdf_document[page_number]
            return page.get_text()
        except Exception as e:
            return f"Error extracting text: {e}"
    
    def close_document(self):
        """Close PDF document."""
        if self.pdf_document:
            self.pdf_document.close()
        super().close_document()
    
    def _create_error_page(self, message: str):
        """Create an error page image."""
        Image, ImageDraw, ImageFont = self._ensure_pil()
        img = Image.new('RGB', (600, 800), color='white')
        draw = ImageDraw.Draw(img)
        draw.text((20, 20), message, fill='red')
        return img


class WordBackend(DocumentBackend):
    """
    Backend for Word document viewing.
    Requires: python-docx library
    
    Note: Uses lazy import - python-docx is only imported when load_document is called.
    """
    
    def __init__(self):
        super().__init__()
        self.doc = None
        self._Document = None
    
    def _ensure_docx(self):
        """Lazy import of python-docx library with auto-install."""
        if self._Document is None:
            if install_library("python-docx", "docx"):
                try:
                    from docx import Document
                    self._Document = Document
                except ImportError:
                    raise ImportError(
                        "python-docx is required for Word support.\n"
                        "Install it with: pip install python-docx"
                    )
            else:
                raise ImportError(
                    "python-docx is required for Word support.\n"
                    "Install it with: pip install python-docx"
                )
        return self._Document
    
    def load_document(self, file_path: str) -> bool:
        """Load Word document."""
        try:
            Document = self._ensure_docx()
            self.doc = Document(file_path)
            self.file_path = file_path
            # For Word, treat each section or page break as a page
            # For simplicity, we'll treat the whole document as pages based on paragraphs
            self.page_count = max(1, (len(self.doc.paragraphs) + 19) // 20)  # Approximate
            self.current_page = 0
            self.document_loaded = True
            return True
        except ImportError as e:
            raise e
        except Exception as e:
            print(f"Error loading Word document: {e}")
            return False
    
    def get_page_count(self) -> int:
        """Get approximate page count."""
        return self.page_count
    
    def render_page(self, page_number: int, width: int = None, height: int = None):
        """Render Word page as image (text-based rendering)."""
        if not self.document_loaded:
            return self._create_error_page("Document not loaded")
        
        Image, ImageDraw, ImageFont = self._ensure_pil()
        
        # Create image with text
        img_width = width or 600
        img_height = height or 800
        img = Image.new('RGB', (img_width, img_height), color='white')
        draw = ImageDraw.Draw(img)
        
        try:
            font = ImageFont.truetype("arial.ttf", 12)
        except:
            font = ImageFont.load_default()
        
        # Get paragraphs for this "page"
        start_para = page_number * 20
        end_para = min(start_para + 20, len(self.doc.paragraphs))
        
        y = 20
        for i in range(start_para, end_para):
            if i < len(self.doc.paragraphs):
                text = self.doc.paragraphs[i].text[:80]  # Limit line length
                if text.strip():
                    draw.text((20, y), text, fill='black', font=font)
                    y += 20
                    if y > img_height - 40:
                        break
        
        return img
    
    def get_page_text(self, page_number: int) -> str:
        """Extract text from Word page."""
        if not self.document_loaded:
            return ""
        
        start_para = page_number * 20
        end_para = min(start_para + 20, len(self.doc.paragraphs))
        
        text_lines = []
        for i in range(start_para, end_para):
            if i < len(self.doc.paragraphs):
                text_lines.append(self.doc.paragraphs[i].text)
        
        return '\n'.join(text_lines)
    
    def close_document(self):
        """Close Word document."""
        self.doc = None
        super().close_document()
    
    def _create_error_page(self, message: str):
        """Create an error page image."""
        Image, ImageDraw, ImageFont = self._ensure_pil()
        img = Image.new('RGB', (600, 800), color='white')
        draw = ImageDraw.Draw(img)
        draw.text((20, 20), message, fill='red')
        return img


class ImageBackend(DocumentBackend):
    """
    Backend for viewing images as documents.
    Supports: JPG, PNG, GIF, BMP, etc.
    """
    
    def __init__(self):
        super().__init__()
        self.image = None
    
    def load_document(self, file_path: str) -> bool:
        """Load image file."""
        try:
            Image, ImageDraw, ImageFont = self._ensure_pil()
            self.image = Image.open(file_path)
            self.file_path = file_path
            self.page_count = 1
            self.current_page = 0
            self.document_loaded = True
            return True
        except Exception as e:
            print(f"Error loading image: {e}")
            return False
    
    def get_page_count(self) -> int:
        """Images have only 1 page."""
        return 1
    
    def render_page(self, page_number: int, width: int = None, height: int = None):
        """Render (resize) image."""
        if not self.document_loaded or page_number != 0:
            return self._create_error_page("Invalid page")
        
        Image, ImageDraw, ImageFont = self._ensure_pil()
        
        if width or height:
            # Resize maintaining aspect ratio
            img_copy = self.image.copy()
            img_copy.thumbnail((width, height), Image.Resampling.LANCZOS)
            return img_copy
        
        return self.image.copy()
    
    def get_page_text(self, page_number: int) -> str:
        """No text in images."""
        return f"Image: {self.file_path}\nSize: {self.image.width}x{self.image.height}"
    
    def close_document(self):
        """Close image."""
        if self.image:
            self.image.close()
        super().close_document()
    
    def _create_error_page(self, message: str):
        """Create an error page image."""
        Image, ImageDraw, ImageFont = self._ensure_pil()
        img = Image.new('RGB', (600, 800), color='white')
        draw = ImageDraw.Draw(img)
        draw.text((20, 20), message, fill='red')
        return img


class TextBackend(DocumentBackend):
    """
    Backend for viewing text files as documents.
    Supports: TXT, MD, LOG, etc.
    """
    
    def __init__(self, lines_per_page: int = 40):
        super().__init__()
        self.lines = []
        self.lines_per_page = lines_per_page
    
    def load_document(self, file_path: str) -> bool:
        """Load text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                self.lines = f.readlines()
            
            self.file_path = file_path
            self.page_count = max(1, (len(self.lines) + self.lines_per_page - 1) // self.lines_per_page)
            self.current_page = 0
            self.document_loaded = True
            return True
        except Exception as e:
            print(f"Error loading text file: {e}")
            return False
    
    def get_page_count(self) -> int:
        """Get page count."""
        return self.page_count
    
    def render_page(self, page_number: int, width: int = None, height: int = None):
        """Render text page as image."""
        if not self.document_loaded or page_number >= self.page_count:
            return self._create_error_page("Invalid page")
        
        Image, ImageDraw, ImageFont = self._ensure_pil()
        
        img_width = width or 700
        img_height = height or 900
        img = Image.new('RGB', (img_width, img_height), color='white')
        draw = ImageDraw.Draw(img)
        
        try:
            font = ImageFont.truetype("consola.ttf", 11)
        except:
            font = ImageFont.load_default()
        
        # Get lines for this page
        start_line = page_number * self.lines_per_page
        end_line = min(start_line + self.lines_per_page, len(self.lines))
        
        y = 20
        for i in range(start_line, end_line):
            line = self.lines[i].rstrip()[:100]  # Limit line length
            draw.text((20, y), line, fill='black', font=font)
            y += 15
        
        return img
    
    def get_page_text(self, page_number: int) -> str:
        """Get text from page."""
        if not self.document_loaded or page_number >= self.page_count:
            return ""
        
        start_line = page_number * self.lines_per_page
        end_line = min(start_line + self.lines_per_page, len(self.lines))
        
        return ''.join(self.lines[start_line:end_line])
    
    def close_document(self):
        """Close text document."""
        self.lines = []
        super().close_document()
    
    def _create_error_page(self, message: str):
        """Create an error page image."""
        Image, ImageDraw, ImageFont = self._ensure_pil()
        img = Image.new('RGB', (600, 800), color='white')
        draw = ImageDraw.Draw(img)
        draw.text((20, 20), message, fill='red')
        return img
